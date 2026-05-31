"""Document parsers for the water review pipeline."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from app.services.water_review_models import DEFAULT_MINERU_JSON, DEFAULT_MINERU_MD, ParsedBlock
from app.services.water_review_utils import (
    _caption_from_text,
    _classify_block,
    _heading_level,
    _html_to_text,
    _number_list,
    _optional_int,
    _replace_section_level,
    _section_hint,
    _section_path,
    _unique_strings,
)

def parse_document(file_path: str | None = None) -> list[ParsedBlock]:
    """Load an explicit source first, then fall back to bundled MinerU samples."""
    if file_path:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == ".json" and path.exists():
            return _parse_mineru_json(path)
        if suffix == ".md" and path.exists():
            return _parse_markdown(path)
        if suffix == ".pdf":
            return _parse_pdf(str(path))
        if suffix == ".docx":
            return _parse_docx(str(path))
        return []

    if DEFAULT_MINERU_JSON.exists():
        return _parse_mineru_json(DEFAULT_MINERU_JSON)
    if DEFAULT_MINERU_MD.exists():
        return _parse_markdown(DEFAULT_MINERU_MD)
    return []

def _parse_mineru_json(path: Path) -> list[ParsedBlock]:
    data = json.loads(path.read_text(encoding="utf-8"))
    pages = data.get("pdf_info", []) if isinstance(data, dict) else []
    blocks: list[ParsedBlock] = []
    cursor = 0
    section_stack: list[tuple[int, str]] = []

    for page in pages:
        page_num = int(page.get("page_idx", 0)) + 1
        page_size = _number_list(page.get("page_size", []))
        for raw in page.get("para_blocks", []) or []:
            html, image_path = _mineru_block_media(raw)
            text = _mineru_block_text(raw).strip()
            if not text and html:
                text = _html_to_text(html)
            if not text and image_path:
                text = image_path
            if not text and not html and not image_path:
                continue
            block_type = _mineru_type(raw.get("type", "text"), html=html, image_path=image_path)
            start = cursor
            cursor += len(text) + 1
            block_index = raw.get("index", len(blocks))
            structure = _mineru_block_structure(raw)
            if block_type == "title" and _is_section_title(text):
                section_stack = _update_section_stack(section_stack, text)
            parent_section = _section_path(section_stack) or _section_hint(text)
            blocks.append(
                ParsedBlock(
                    block_id=f"p{page_num}-b{block_index}",
                    page=page_num,
                    bbox=_number_list(raw.get("bbox", [])),
                    text=text,
                    type=block_type,
                    section_hint=parent_section,
                    char_start=start,
                    char_end=start + len(text),
                    html=html,
                    image_path=image_path,
                    mineru_index=_optional_int(block_index),
                    mineru_sub_type=str(raw.get("sub_type") or ""),
                    page_size=page_size,
                    child_types=structure["child_types"],
                    span_types=structure["span_types"],
                    line_bboxes=structure["line_bboxes"],
                    span_bboxes=structure["span_bboxes"],
                    parent_section=parent_section,
                    caption=_mineru_caption(raw, text),
                    atomic_index=len(blocks),
                )
            )
    if not blocks and DEFAULT_MINERU_MD.exists():
        blocks = _parse_markdown(DEFAULT_MINERU_MD)
    return blocks


def _parse_markdown(path: Path) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    cursor = 0
    section_stack: list[tuple[int, str]] = []
    for index, raw_line in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
        text = raw_line.strip()
        if not text:
            continue
        block_type = "title" if text.startswith("#") else _classify_block(text)
        clean_text = text.lstrip("#").strip()
        if block_type == "title":
            level = max(1, min(6, len(text) - len(text.lstrip("#"))))
            section_stack = _replace_section_level(section_stack, level, clean_text)
        parent_section = _section_path(section_stack) or _section_hint(clean_text)
        start = cursor
        cursor += len(clean_text) + 1
        blocks.append(
            ParsedBlock(
                block_id=f"md-b{index:05d}",
                page=max(1, index // 40 + 1),
                bbox=[],
                text=clean_text,
                type=block_type,
                section_hint=parent_section,
                char_start=start,
                char_end=start + len(clean_text),
                parent_section=parent_section,
                caption=_caption_from_text(clean_text),
                atomic_index=len(blocks),
            )
        )
    return blocks


def _mineru_type(raw_type: str, html: str = "", image_path: str = "") -> str:
    if raw_type == "title":
        return "title"
    if raw_type == "table":
        return "table"
    if raw_type == "image":
        return "image"
    if html:
        return "table"
    if image_path:
        return "image"
    return "paragraph"


def _mineru_block_text(block: dict[str, Any]) -> str:
    parts: list[str] = []
    if "lines" in block:
        parts.extend(_mineru_lines_text(block.get("lines", [])))
    for child in block.get("blocks", []) or []:
        parts.extend(_mineru_lines_text(child.get("lines", [])))
    return "\n".join(part for part in parts if part.strip())


def _mineru_caption(block: dict[str, Any], fallback_text: str) -> str:
    captions: list[str] = []
    for child in block.get("blocks", []) or []:
        child_type = str(child.get("type") or "")
        if "caption" in child_type:
            captions.extend(_mineru_lines_text(child.get("lines", [])))
    return _caption_from_text("\n".join(captions) or fallback_text)


def _mineru_lines_text(lines: list[dict[str, Any]]) -> list[str]:
    result: list[str] = []
    for line in lines:
        spans = line.get("spans", []) or []
        text = "".join(str(span.get("content", "")) for span in spans).strip()
        if text:
            result.append(text)
    return result


def _mineru_block_media(block: dict[str, Any]) -> tuple[str, str]:
    html = str(block.get("html") or "").strip()
    image_path = str(block.get("image_path") or "").strip()
    for line in block.get("lines", []) or []:
        for span in line.get("spans", []) or []:
            if not html and span.get("html"):
                html = str(span.get("html") or "").strip()
            if not image_path and span.get("image_path"):
                image_path = str(span.get("image_path") or "").strip()
    for child in block.get("blocks", []) or []:
        child_html, child_image_path = _mineru_block_media(child)
        if not html:
            html = child_html
        if not image_path:
            image_path = child_image_path
    return html, image_path


def _mineru_block_structure(block: dict[str, Any]) -> dict[str, list[Any]]:
    child_types: list[str] = []
    span_types: list[str] = []
    line_bboxes: list[list[float]] = []
    span_bboxes: list[list[float]] = []

    def visit(raw: dict[str, Any], include_self: bool = False) -> None:
        raw_type = str(raw.get("type") or "").strip()
        if include_self and raw_type:
            child_types.append(raw_type)
        for line in raw.get("lines", []) or []:
            line_bbox = _number_list(line.get("bbox", []))
            if line_bbox:
                line_bboxes.append(line_bbox)
            for span in line.get("spans", []) or []:
                span_type = str(span.get("type") or "").strip()
                if span_type:
                    span_types.append(span_type)
                span_bbox = _number_list(span.get("bbox", []))
                if span_bbox:
                    span_bboxes.append(span_bbox)
        for child in raw.get("blocks", []) or []:
            visit(child, include_self=True)

    visit(block)
    return {
        "child_types": _unique_strings(child_types),
        "span_types": _unique_strings(span_types),
        "line_bboxes": line_bboxes[:20],
        "span_bboxes": span_bboxes[:40],
    }

def _update_section_stack(section_stack: list[tuple[int, str]], title: str) -> list[tuple[int, str]]:
    return _replace_section_level(section_stack, _heading_level(title), title.strip())


def _is_section_title(title: str) -> bool:
    stripped = title.strip()
    if not stripped or len(stripped) > 80:
        return False
    if _section_hint(stripped):
        return True
    if re.match(r"^第[一二三四五六七八九十\d]+章", stripped):
        return True
    if re.match(r"^\d+(?:\.\d+)*\s+\S", stripped):
        return True
    if re.match(r"^[一二三四五六七八九十]+[、.．]", stripped):
        return True
    if re.match(r"^[（(][一二三四五六七八九十\d]+[）)]", stripped):
        return True
    return False


def _parse_pdf(file_path: str) -> list[ParsedBlock]:
    import fitz

    blocks: list[ParsedBlock] = []
    cursor = 0
    with fitz.open(file_path) as doc:
        for page_index, page in enumerate(doc, start=1):
            data = page.get_text("dict")
            for raw in data.get("blocks", []):
                lines = []
                for line in raw.get("lines", []):
                    line_text = "".join(
                        span.get("text", "")
                        for span in line.get("spans", [])
                    ).strip()
                    if line_text:
                        lines.append(line_text)
                text = "\n".join(lines).strip()
                if not text:
                    continue
                start = cursor
                cursor += len(text) + 1
                blocks.append(
                    ParsedBlock(
                        block_id=f"p{page_index}-b{len(blocks) + 1:05d}",
                        page=page_index,
                        bbox=[round(float(v), 2) for v in raw.get("bbox", [])],
                        text=text,
                        type=_classify_block(text),
                        section_hint=_section_hint(text),
                        char_start=start,
                        char_end=start + len(text),
                    )
                )
    return blocks


def _parse_docx(file_path: str) -> list[ParsedBlock]:
    from docx import Document

    blocks: list[ParsedBlock] = []
    cursor = 0
    doc = Document(file_path)
    for index, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if not text:
            continue
        start = cursor
        cursor += len(text) + 1
        blocks.append(
            ParsedBlock(
                block_id=f"docx-b{index:05d}",
                page=1,
                bbox=[],
                text=text,
                type=_classify_block(text),
                section_hint=_section_hint(text),
                char_start=start,
                char_end=start + len(text),
            )
        )
    for table_index, table in enumerate(doc.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        text = "\n".join(row for row in rows if row.strip())
        if not text:
            continue
        start = cursor
        cursor += len(text) + 1
        blocks.append(
            ParsedBlock(
                block_id=f"docx-t{table_index:05d}",
                page=1,
                bbox=[],
                text=text,
                type="table",
                section_hint="表格",
                char_start=start,
                char_end=start + len(text),
            )
        )
    return blocks
