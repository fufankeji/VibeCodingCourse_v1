"""Water-soil conservation review pipeline primitives.

This module implements the first-version domain path:
- coordinate-aware MinerU JSON consumption
- PyMuPDF / DOCX fallback parsing when prepared MinerU artifacts are absent
- bbox-aware chunk building
- deterministic field extraction with source spans
- Chroma/SiliconFlow/DeepSeek RAG issue generation

The public functions intentionally return plain dicts so they can feed the
existing LangGraph/ReviewItem MVP without a database migration.
"""

from __future__ import annotations

import json
import re
import uuid
from dataclasses import asdict, dataclass, field
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

from app.config import settings
from app.services.mineru_table_fact_service import extract_table_facts
from app.services.review_rule_schema import build_review_rule_topics, normalize_review_rules


@dataclass
class ParsedBlock:
    block_id: str
    page: int
    bbox: list[float]
    text: str
    type: str
    section_hint: str
    char_start: int
    char_end: int
    html: str = ""
    image_path: str = ""
    mineru_index: int | None = None
    mineru_sub_type: str = ""
    page_size: list[float] = field(default_factory=list)
    child_types: list[str] = field(default_factory=list)
    span_types: list[str] = field(default_factory=list)
    line_bboxes: list[list[float]] = field(default_factory=list)
    span_bboxes: list[list[float]] = field(default_factory=list)
    parent_section: str = ""
    caption: str = ""
    atomic_index: int = 0


@dataclass
class ReviewChunk:
    chunk_id: str
    text: str
    section: str
    page_range: list[int]
    bbox_list: list[dict[str, Any]]
    table_refs: list[str]
    metadata: dict[str, Any]
    char_start: int
    char_end: int
    embedding_text: str = ""


WATER_FIELDS = [
    "project_name",
    "construction_unit",
    "construction_location",
    "project_nature",
    "key_prevention_or_control_area",
    "disturbed_area",
    "land_area",
    "prevention_responsibility_area",
    "zone_area",
    "excavation_volume",
    "fill_volume",
    "borrow_volume",
    "spoil_volume",
    "comprehensive_utilization",
    "spoil_destination",
    "topsoil_stripping",
    "topsoil_preservation",
    "topsoil_backfill",
    "temp_soil_stockpile",
    "borrow_area",
    "spoil_area",
    "construction_road",
    "prevention_measures",
    "monitoring",
    "schedule_arrangement",
    "investment_estimate",
]


SECTION_KEYWORDS = {
    "项目概况": ["项目概况", "综合说明", "工程概况"],
    "防治责任范围": ["防治责任范围", "扰动范围", "占地"],
    "水土流失分析": ["水土流失", "流失预测", "流失现状"],
    "防治措施": ["水土保持措施", "防治措施", "工程措施", "植物措施"],
    "监测": ["水土保持监测", "监测"],
    "投资估算": ["投资估算", "效益分析", "投资"],
    "结论": ["结论", "建议"],
}


RULES = [
    {
        "rule_id": "SWC-FORM-001",
        "rule_name": "必备章节完整性检查",
        "rule_category": "形式类",
        "severity": "HIGH",
        "expected": "报告应包含项目概况、防治责任范围、防治措施、监测、投资估算、结论等核心章节",
    },
    {
        "rule_id": "SWC-FIELD-001",
        "rule_name": "项目基础信息完整性检查",
        "rule_category": "形式类",
        "severity": "MEDIUM",
        "expected": "项目名称、建设单位、建设地点、建设性质应有明确表述",
    },
    {
        "rule_id": "SWC-TECH-001",
        "rule_name": "扰动面积明确性检查",
        "rule_category": "技术类",
        "severity": "HIGH",
        "expected": "应明确扰动地表面积或防治责任范围面积",
    },
    {
        "rule_id": "SWC-TECH-002",
        "rule_name": "土石方平衡一致性检查",
        "rule_category": "一致性类",
        "severity": "HIGH",
        "expected": "挖方 + 借方 应与 填方 + 弃方 基本平衡",
    },
    {
        "rule_id": "SWC-TECH-003",
        "rule_name": "表土保护措施检查",
        "rule_category": "技术类",
        "severity": "MEDIUM",
        "expected": "涉及扰动地表时应说明表土剥离、保存或回覆措施",
    },
    {
        "rule_id": "SWC-TECH-004",
        "rule_name": "弃方与弃渣场匹配检查",
        "rule_category": "技术类",
        "severity": "HIGH",
        "expected": "存在弃方时应说明弃渣场或弃土去向",
    },
    {
        "rule_id": "SWC-ATTR-001",
        "rule_name": "区域属性明确性检查",
        "rule_category": "项目属性判定",
        "severity": "LOW",
        "expected": "应说明是否位于重点预防区、重点治理区或易发生水土流失区域",
    },
]

BACKEND_DIR = Path(__file__).resolve().parents[2]
DATA_DIR = BACKEND_DIR / "data"
DEFAULT_MINERU_JSON = DATA_DIR / "MinerU_1 北京航空航天大学沙河校区图书馆项目水土保持方案.json"
DEFAULT_MINERU_MD = DATA_DIR / "北京航空航天大学沙河校区图书馆项目-mineru.md"
DEFAULT_RULE_SET = DATA_DIR / "水土保持方案审查规则集.json"
QUICK_VALIDATION_ISSUE_COUNT = 10
SEMANTIC_CHUNK_MAX_CHARS = 1100
SEMANTIC_CHUNK_OVERLAP_BLOCKS = 1
EVIDENCE_WINDOW_BLOCK_RADIUS = 1
TABLE_ROW_MAX_ROWS_PER_TABLE = 80
TABLE_ROW_INCLUDE_KEYWORDS = (
    "弃渣",
    "弃土",
    "弃方",
    "取土",
    "取料",
    "土石方",
    "表土",
    "占地",
    "防治责任",
    "防治措施",
    "监测",
    "扰动",
)
TABLE_ROW_EXCLUDE_HINTS = ("工程单价", "机械台时", "概算附表", "单价汇总", "附表")


def run_pipeline(file_path: str, artifact_dir: str, session_id: str) -> dict[str, Any]:
    """Parse, chunk, extract fields, review rules, and persist JSON artifacts."""
    blocks = parse_document(file_path)
    chunks = build_chunks(blocks)
    fallback_fields = extract_fields(chunks)
    table_facts = extract_table_facts(blocks, chunks)
    langextract_facts: list[dict[str, Any]] = list(table_facts)
    cross_chapter_findings: list[dict[str, Any]] = []
    if settings.langextract_enabled:
        from app.services.langextract_service import (
            build_cross_chapter_findings,
            build_fact_index,
            facts_to_extracted_fields,
            run_langextract,
        )

        langextract_facts = [*table_facts, *run_langextract(chunks)]
        fields = facts_to_extracted_fields(langextract_facts, fallback_fields)
        fact_index = build_fact_index(langextract_facts)
        cross_chapter_findings = build_cross_chapter_findings(langextract_facts)
    else:
        if table_facts:
            from app.services.langextract_service import build_fact_index, facts_to_extracted_fields

            fields = facts_to_extracted_fields(table_facts, fallback_fields)
            fact_index = build_fact_index(table_facts)
        else:
            fields = fallback_fields
            fact_index = {"fact_count": 0, "fields": [], "by_field": {}}
    rules = load_rule_set()

    artifact_path = Path(artifact_dir)
    artifact_path.mkdir(parents=True, exist_ok=True)
    from app.services.rag_service import run_rag_review

    rag_result = run_rag_review(
        session_id,
        chunks,
        rules,
        artifact_path,
        facts=langextract_facts,
        findings=cross_chapter_findings,
    )
    from app.services.review_config_service import list_check_item_specs

    configured_check_items = [item for item in list_check_item_specs() if item.get("enabled") is not False]
    configured_issues = _issues_from_configured_rules(
        session_id,
        chunks,
        "\n".join(chunk.text for chunk in chunks),
        fields,
        configured_check_items,
        limit=len(configured_check_items),
        evidence_store=_pipeline_evidence_store(session_id, rag_result) if configured_check_items else None,
    )
    issues = [*rag_result["issues"], *configured_issues]
    rule_topics = build_review_rule_topics(rules, issues, configured_check_items=configured_check_items)
    _write_json(artifact_path / "parsed_blocks.json", [asdict(b) for b in blocks])
    _write_json(artifact_path / "review_chunks.json", [asdict(c) for c in chunks])
    _write_json(artifact_path / "extracted_fields.json", fields)
    _write_json(artifact_path / "langextract_facts.json", langextract_facts)
    _write_json(artifact_path / "langextract_fact_index.json", fact_index)
    _write_json(artifact_path / "cross_chapter_findings.json", cross_chapter_findings)
    _write_json(artifact_path / "review_rule_topics.json", rule_topics)
    _write_json(artifact_path / "issues.json", issues)

    return {
        "full_text": "\n".join(block.text for block in blocks),
        "blocks": blocks,
        "chunks": chunks,
        "fields": fields,
        "facts": langextract_facts,
        "cross_chapter_findings": cross_chapter_findings,
        "review_items": issues,
        "rules": rules,
        "rule_topics": rule_topics,
        "rag": rag_result,
        "artifact_dir": str(artifact_path),
    }


def _pipeline_evidence_store(session_id: str, rag_result: dict[str, Any]) -> Any | None:
    manifest = rag_result.get("index_manifest") if isinstance(rag_result, dict) else {}
    if not isinstance(manifest, dict):
        return None
    vector_store = str(manifest.get("vector_store") or "").strip()
    if not vector_store:
        return None
    try:
        from app.services.rag_service import ChromaChunkStore, SiliconFlowEmbeddingProvider

        return ChromaChunkStore(Path(vector_store), session_id, SiliconFlowEmbeddingProvider())
    except Exception:
        return None


def parse_document(file_path: str | None = None) -> list[ParsedBlock]:
    """Load an explicit source first, then fall back to bundled MinerU samples."""
    if file_path:
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == ".json" and path.exists():
            return _parse_mineru_json(path)
        if suffix == ".md" and path.exists():
            return _parse_markdown(path)
        if suffix == ".pdf":
            return _parse_pdf(str(path))
        if suffix == ".docx":
            return _parse_docx(str(path))
        return []

    if DEFAULT_MINERU_JSON.exists():
        return _parse_mineru_json(DEFAULT_MINERU_JSON)
    if DEFAULT_MINERU_MD.exists():
        return _parse_markdown(DEFAULT_MINERU_MD)
    return []


def load_rule_set(path: Path = DEFAULT_RULE_SET) -> list[dict[str, Any]]:
    if not path.exists():
        return normalize_review_rules(RULES)
    data = json.loads(path.read_text(encoding="utf-8"))
    rules = data.get("rules", []) if isinstance(data, dict) else data
    return normalize_review_rules(rules or RULES)


def build_chunks(blocks: list[ParsedBlock], max_chars: int = SEMANTIC_CHUNK_MAX_CHARS) -> list[ReviewChunk]:
    chunks: list[ReviewChunk] = []
    current: list[ParsedBlock] = []
    block_positions = {block.block_id: index for index, block in enumerate(blocks)}

    def next_chunk_id() -> str:
        return f"chunk-{len(chunks) + 1:04d}"

    def flush() -> None:
        nonlocal current
        core_blocks = [block for block in current if _block_document_text(block)]
        if not core_blocks:
            current = []
            return
        chunk = _semantic_chunk_from_blocks(next_chunk_id(), core_blocks, blocks, block_positions)
        if not chunk.text.strip():
            current = []
            return
        chunks.append(chunk)
        for table_block in [block for block in core_blocks if block.html and _should_emit_table_row_chunks(block)]:
            for row_text in _table_row_texts(table_block):
                chunks.append(_table_row_chunk(next_chunk_id(), table_block, row_text, chunk.chunk_id))
        current = []

    for block in blocks:
        if block.type == "title":
            if any(item.type != "title" and _block_document_text(item) for item in current):
                flush()
            current.append(block)
            continue
        projected_len = sum(_block_semantic_size(b) for b in current) + _block_semantic_size(block)
        if current and _block_section(current[-1]) != _block_section(block):
            flush()
        elif current and projected_len > max_chars:
            overlap = _overlap_blocks(current)
            flush()
            current.extend(overlap)
        current.append(block)
    flush()
    return chunks


def extract_fields(chunks: list[ReviewChunk]) -> list[dict[str, Any]]:
    text = "\n".join(chunk.text for chunk in chunks)
    extracted = [
        _field("project_name", _match_after(text, [r"项目名称[:：]?\s*([^\n，。；;]+)"]), chunks),
        _field("construction_unit", _match_after(text, [r"建设单位[:：]?\s*([^\n，。；;]+)", r"建设方[:：]?\s*([^\n，。；;]+)"]), chunks),
        _field("construction_location", _match_after(text, [r"建设地点[:：]?\s*([^\n，。；;]+)", r"项目地点[:：]?\s*([^\n，。；;]+)"]), chunks),
        _field("project_nature", _match_after(text, [r"建设性质[:：]?\s*([^\n，。；;]+)", r"项目性质[:：]?\s*([^\n，。；;]+)"]), chunks),
        _field("disturbed_area", _match_metric(text, ["扰动地表面积", "扰动面积", "防治责任范围"]), chunks),
        _field("land_area", _match_metric(text, ["占地面积", "总占地"]), chunks),
        _field("prevention_responsibility_area", _match_metric(text, ["防治责任范围面积", "防治责任范围"]), chunks),
        _field("zone_area", _match_metric(text, ["分区面积", "防治分区"]), chunks),
        _field("excavation_volume", _match_metric(text, ["挖方", "开挖量"]), chunks),
        _field("fill_volume", _match_metric(text, ["填方", "回填量"]), chunks),
        _field("borrow_volume", _match_metric(text, ["借方", "取土量"]), chunks),
        _field("spoil_volume", _match_metric(text, ["弃方", "弃渣量", "弃土量"]), chunks),
        _field("investment_estimate", _match_metric(text, ["投资估算", "水土保持投资", "总投资"]), chunks),
    ]

    keyword_fields = {
        "key_prevention_or_control_area": ["重点预防区", "重点治理区", "易发生水土流失"],
        "topsoil_stripping": ["表土剥离"],
        "topsoil_preservation": ["表土保存", "表土临时堆存", "表土堆存"],
        "topsoil_backfill": ["表土回覆", "表土回填"],
        "comprehensive_utilization": ["综合利用"],
        "spoil_destination": ["外运", "消纳场", "弃土去向", "弃方去向"],
        "temp_soil_stockpile": ["临时堆土区", "临时堆土场"],
        "borrow_area": ["取土场"],
        "spoil_area": ["弃渣场", "弃土场"],
        "construction_road": ["施工道路", "施工便道"],
        "prevention_measures": ["工程措施", "植物措施", "临时措施", "防治措施"],
        "monitoring": ["水土保持监测", "监测"],
        "schedule_arrangement": ["时序安排", "施工时序", "实施进度"],
    }
    for name, keywords in keyword_fields.items():
        extracted.append(_field(name, _match_keyword(text, keywords), chunks))

    present = {item["field_name"]: item for item in extracted}
    return [present.get(name) or _field(name, None, chunks) for name in WATER_FIELDS]


def review_rules(
    session_id: str,
    chunks: list[ReviewChunk],
    fields: list[dict[str, Any]],
    rules: list[dict[str, Any]] | None = None,
    evidence_store: Any | None = None,
) -> list[dict[str, Any]]:
    text = "\n".join(chunk.text for chunk in chunks)
    by_name = {field["field_name"]: field for field in fields}
    issues: list[dict[str, Any]] = []
    configured_rules = rules or RULES

    found_sections = _detect_sections(text)
    missing_sections = [name for name in ["项目概况", "防治责任范围", "防治措施", "监测", "投资估算", "结论"] if name not in found_sections]
    if missing_sections:
        issues.append(_issue(session_id, RULES[0], "缺少或未能识别必备章节：" + "、".join(missing_sections), "", "已识别章节：" + "、".join(found_sections), RULES[0]["expected"], _best_chunk(chunks, ["目录", "项目概况"])))

    issues.extend(
        _issues_from_configured_rules(
            session_id,
            chunks,
            text,
            fields,
            configured_rules,
            limit=QUICK_VALIDATION_ISSUE_COUNT - len(issues),
            evidence_store=evidence_store,
        )
    )

    missing_basic = [
        label
        for key, label in [
            ("project_name", "项目名称"),
            ("construction_unit", "建设单位"),
            ("construction_location", "建设地点"),
            ("project_nature", "建设性质"),
        ]
        if not by_name[key]["value"]
    ]
    if missing_basic:
        _append_if_room(issues, _issue(session_id, RULES[1], "基础信息缺失或未见明确表述：" + "、".join(missing_basic), "", "缺失字段：" + "、".join(missing_basic), RULES[1]["expected"], _best_chunk(chunks, ["项目概况", "综合说明"])))

    if not by_name["disturbed_area"]["value"]:
        _append_if_room(issues, _issue(session_id, RULES[2], "未见扰动地表面积或防治责任范围面积的明确数值。", "", "材料中未见明确表述", RULES[2]["expected"], _best_chunk(chunks, ["扰动", "防治责任范围", "占地"])))

    excavation = _to_number(by_name["excavation_volume"]["normalized_value"])
    fill = _to_number(by_name["fill_volume"]["normalized_value"])
    borrow = _to_number(by_name["borrow_volume"]["normalized_value"])
    spoil = _to_number(by_name["spoil_volume"]["normalized_value"])
    if all(value is not None for value in [excavation, fill, borrow, spoil]):
        left = excavation + borrow
        right = fill + spoil
        tolerance = max(abs(left), abs(right), 1.0) * 0.05
        if abs(left - right) > tolerance:
            actual = f"挖方+借方={left:g}，填方+弃方={right:g}"
            _append_if_room(issues, _issue(session_id, RULES[3], "土石方平衡关系存在不一致。", "", actual, RULES[3]["expected"], _best_chunk(chunks, ["土石方", "挖方", "填方", "弃方"])))
    elif any(value is not None for value in [excavation, fill, borrow, spoil]):
        _append_if_room(issues, _issue(session_id, RULES[3], "土石方平衡关键字段不完整，无法完成一致性核验。", "", "已抽取部分土石方字段，但缺少挖/填/借/弃中的一项或多项", RULES[3]["expected"], _best_chunk(chunks, ["土石方", "挖方", "填方", "弃方"]), risk_override="MEDIUM"))

    has_disturbance = bool(by_name["disturbed_area"]["value"] or by_name["land_area"]["value"])
    has_topsoil = any(by_name[key]["value"] for key in ["topsoil_stripping", "topsoil_preservation", "topsoil_backfill"])
    if has_disturbance and not has_topsoil:
        _append_if_room(issues, _issue(session_id, RULES[4], "存在扰动或占地信息，但未见表土剥离、保存或回覆措施。", "", "材料中未见明确表土保护措施", RULES[4]["expected"], _best_chunk(chunks, ["表土", "防治措施", "扰动"])))

    if spoil and spoil > 0 and not by_name["spoil_area"]["value"]:
        _append_if_room(issues, _issue(session_id, RULES[5], "存在弃方/弃渣量，但未见弃渣场或弃土去向说明。", "", f"弃方/弃渣量={spoil:g}", RULES[5]["expected"], _best_chunk(chunks, ["弃方", "弃渣", "弃土"])))

    if not by_name["key_prevention_or_control_area"]["value"]:
        _append_if_room(issues, _issue(session_id, RULES[6], "未见项目所在区域属性说明。", "", "材料中未见重点预防区、重点治理区或易发生水土流失区域表述", RULES[6]["expected"], _best_chunk(chunks, ["区域", "水土流失", "项目区"])))

    if not issues:
        summary_chunk = chunks[0] if chunks else None
        issues.append(_issue(session_id, RULES[6], "首版规则未发现高/中风险问题，建议人工抽查关键字段和图表一致性。", "", "自动规则未命中", "人工复核确认", summary_chunk, risk_override="LOW", confidence=55))

    return issues[:QUICK_VALIDATION_ISSUE_COUNT]


def _append_if_room(issues: list[dict[str, Any]], issue: dict[str, Any]) -> None:
    if len(issues) < QUICK_VALIDATION_ISSUE_COUNT:
        issues.append(issue)


def _parse_mineru_json(path: Path) -> list[ParsedBlock]:
    data = json.loads(path.read_text(encoding="utf-8"))
    pages = data.get("pdf_info", []) if isinstance(data, dict) else []
    blocks: list[ParsedBlock] = []
    cursor = 0
    section_stack: list[tuple[int, str]] = []

    for page in pages:
        page_num = int(page.get("page_idx", 0)) + 1
        page_size = _number_list(page.get("page_size", []))
        for raw in page.get("para_blocks", []) or []:
            html, image_path = _mineru_block_media(raw)
            text = _mineru_block_text(raw).strip()
            if not text and html:
                text = _html_to_text(html)
            if not text and image_path:
                text = image_path
            if not text and not html and not image_path:
                continue
            block_type = _mineru_type(raw.get("type", "text"), html=html, image_path=image_path)
            start = cursor
            cursor += len(text) + 1
            block_index = raw.get("index", len(blocks))
            structure = _mineru_block_structure(raw)
            if block_type == "title" and _is_section_title(text):
                section_stack = _update_section_stack(section_stack, text)
            parent_section = _section_path(section_stack) or _section_hint(text)
            blocks.append(
                ParsedBlock(
                    block_id=f"p{page_num}-b{block_index}",
                    page=page_num,
                    bbox=_number_list(raw.get("bbox", [])),
                    text=text,
                    type=block_type,
                    section_hint=parent_section,
                    char_start=start,
                    char_end=start + len(text),
                    html=html,
                    image_path=image_path,
                    mineru_index=_optional_int(block_index),
                    mineru_sub_type=str(raw.get("sub_type") or ""),
                    page_size=page_size,
                    child_types=structure["child_types"],
                    span_types=structure["span_types"],
                    line_bboxes=structure["line_bboxes"],
                    span_bboxes=structure["span_bboxes"],
                    parent_section=parent_section,
                    caption=_mineru_caption(raw, text),
                    atomic_index=len(blocks),
                )
            )
    if not blocks and DEFAULT_MINERU_MD.exists():
        blocks = _parse_markdown(DEFAULT_MINERU_MD)
    return blocks


def _parse_markdown(path: Path) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    cursor = 0
    section_stack: list[tuple[int, str]] = []
    for index, raw_line in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
        text = raw_line.strip()
        if not text:
            continue
        block_type = "title" if text.startswith("#") else _classify_block(text)
        clean_text = text.lstrip("#").strip()
        if block_type == "title":
            level = max(1, min(6, len(text) - len(text.lstrip("#"))))
            section_stack = _replace_section_level(section_stack, level, clean_text)
        parent_section = _section_path(section_stack) or _section_hint(clean_text)
        start = cursor
        cursor += len(clean_text) + 1
        blocks.append(
            ParsedBlock(
                block_id=f"md-b{index:05d}",
                page=max(1, index // 40 + 1),
                bbox=[],
                text=clean_text,
                type=block_type,
                section_hint=parent_section,
                char_start=start,
                char_end=start + len(clean_text),
                parent_section=parent_section,
                caption=_caption_from_text(clean_text),
                atomic_index=len(blocks),
            )
        )
    return blocks


def _mineru_type(raw_type: str, html: str = "", image_path: str = "") -> str:
    if raw_type == "title":
        return "title"
    if raw_type == "table":
        return "table"
    if raw_type == "image":
        return "image"
    if html:
        return "table"
    if image_path:
        return "image"
    return "paragraph"


def _mineru_block_text(block: dict[str, Any]) -> str:
    parts: list[str] = []
    if "lines" in block:
        parts.extend(_mineru_lines_text(block.get("lines", [])))
    for child in block.get("blocks", []) or []:
        parts.extend(_mineru_lines_text(child.get("lines", [])))
    return "\n".join(part for part in parts if part.strip())


def _mineru_caption(block: dict[str, Any], fallback_text: str) -> str:
    captions: list[str] = []
    for child in block.get("blocks", []) or []:
        child_type = str(child.get("type") or "")
        if "caption" in child_type:
            captions.extend(_mineru_lines_text(child.get("lines", [])))
    return _caption_from_text("\n".join(captions) or fallback_text)


def _mineru_lines_text(lines: list[dict[str, Any]]) -> list[str]:
    result: list[str] = []
    for line in lines:
        spans = line.get("spans", []) or []
        text = "".join(str(span.get("content", "")) for span in spans).strip()
        if text:
            result.append(text)
    return result


def _mineru_block_media(block: dict[str, Any]) -> tuple[str, str]:
    html = str(block.get("html") or "").strip()
    image_path = str(block.get("image_path") or "").strip()
    for line in block.get("lines", []) or []:
        for span in line.get("spans", []) or []:
            if not html and span.get("html"):
                html = str(span.get("html") or "").strip()
            if not image_path and span.get("image_path"):
                image_path = str(span.get("image_path") or "").strip()
    for child in block.get("blocks", []) or []:
        child_html, child_image_path = _mineru_block_media(child)
        if not html:
            html = child_html
        if not image_path:
            image_path = child_image_path
    return html, image_path


def _mineru_block_structure(block: dict[str, Any]) -> dict[str, list[Any]]:
    child_types: list[str] = []
    span_types: list[str] = []
    line_bboxes: list[list[float]] = []
    span_bboxes: list[list[float]] = []

    def visit(raw: dict[str, Any], include_self: bool = False) -> None:
        raw_type = str(raw.get("type") or "").strip()
        if include_self and raw_type:
            child_types.append(raw_type)
        for line in raw.get("lines", []) or []:
            line_bbox = _number_list(line.get("bbox", []))
            if line_bbox:
                line_bboxes.append(line_bbox)
            for span in line.get("spans", []) or []:
                span_type = str(span.get("type") or "").strip()
                if span_type:
                    span_types.append(span_type)
                span_bbox = _number_list(span.get("bbox", []))
                if span_bbox:
                    span_bboxes.append(span_bbox)
        for child in raw.get("blocks", []) or []:
            visit(child, include_self=True)

    visit(block)
    return {
        "child_types": _unique_strings(child_types),
        "span_types": _unique_strings(span_types),
        "line_bboxes": line_bboxes[:20],
        "span_bboxes": span_bboxes[:40],
    }


def _block_anchor(block: ParsedBlock) -> dict[str, Any]:
    anchor: dict[str, Any] = {
        "block_id": block.block_id,
        "page": block.page,
        "bbox": block.bbox,
    }
    if len(block.page_size) >= 2:
        anchor["page_width"] = block.page_size[0]
        anchor["page_height"] = block.page_size[1]
    if block.mineru_index is not None:
        anchor["mineru_index"] = block.mineru_index
    if block.type:
        anchor["block_type"] = block.type
    if block.parent_section:
        anchor["parent_section"] = block.parent_section
    return anchor


def _semantic_chunk_from_blocks(
    chunk_id: str,
    core_blocks: list[ParsedBlock],
    all_blocks: list[ParsedBlock],
    block_positions: dict[str, int],
) -> ReviewChunk:
    section = _semantic_chunk_section(core_blocks)
    text = _semantic_chunk_display_text(core_blocks, section)
    window_blocks = _evidence_window_blocks(core_blocks, all_blocks, block_positions)
    pages = sorted({block.page for block in core_blocks})
    metadata = _chunk_structure_metadata(core_blocks)
    metadata.update(
        {
            "chunk_layer": "semantic",
            "parent_section": section,
            "atomic_block_ids": [block.block_id for block in core_blocks],
            "evidence_window_block_ids": [block.block_id for block in window_blocks],
            "evidence_window_text": _semantic_chunk_display_text(window_blocks, section),
            "evidence_window_bbox_list": [_block_anchor(block) for block in window_blocks if block.bbox],
        }
    )
    return ReviewChunk(
        chunk_id=chunk_id,
        text=text,
        section=section,
        page_range=[pages[0], pages[-1]] if pages else [1, 1],
        bbox_list=[_block_anchor(block) for block in core_blocks if block.bbox],
        table_refs=[block.block_id for block in core_blocks if block.type in {"table", "cell"} or block.html],
        metadata=metadata,
        char_start=core_blocks[0].char_start,
        char_end=core_blocks[-1].char_end,
        embedding_text=_build_chunk_embedding_text(core_blocks, section, text),
    )


def _table_row_chunk(chunk_id: str, table_block: ParsedBlock, row_text: str, parent_chunk_id: str) -> ReviewChunk:
    section = _block_section(table_block)
    text = "\n".join(_unique_strings([section, table_block.caption, row_text]))
    metadata = _chunk_structure_metadata([table_block])
    metadata.update(
        {
            "chunk_layer": "table_row",
            "parent_section": section,
            "parent_chunk_id": parent_chunk_id,
            "atomic_block_ids": [table_block.block_id],
            "evidence_window_block_ids": [table_block.block_id],
            "evidence_window_text": text,
            "evidence_window_bbox_list": [_block_anchor(table_block)] if table_block.bbox else [],
        }
    )
    return ReviewChunk(
        chunk_id=chunk_id,
        text=text,
        section=section,
        page_range=[table_block.page, table_block.page],
        bbox_list=[_block_anchor(table_block)] if table_block.bbox else [],
        table_refs=[table_block.block_id],
        metadata=metadata,
        char_start=table_block.char_start,
        char_end=table_block.char_end,
        embedding_text=text,
    )


def _chunk_structure_metadata(blocks: list[ParsedBlock]) -> dict[str, Any]:
    page_sizes: dict[str, list[float]] = {}
    for block in blocks:
        if len(block.page_size) >= 2:
            page_sizes[str(block.page)] = block.page_size[:2]
    return {
        "block_ids": [b.block_id for b in blocks],
        "block_types": _unique_strings([b.type for b in blocks]),
        "block_sub_types": _unique_strings([b.mineru_sub_type for b in blocks if b.mineru_sub_type]),
        "mineru_indexes": [b.mineru_index for b in blocks if b.mineru_index is not None],
        "child_types": _unique_strings(item for b in blocks for item in b.child_types),
        "span_types": _unique_strings(item for b in blocks for item in b.span_types),
        "page_sizes": page_sizes,
        "has_table": any(b.type == "table" or bool(b.html) for b in blocks),
        "has_image": any(b.type == "image" or bool(b.image_path) for b in blocks),
        "captions": _unique_strings([b.caption for b in blocks if b.caption]),
    }


def _build_chunk_embedding_text(blocks: list[ParsedBlock], section: str, display_text: str) -> str:
    lines: list[str] = []
    if section != "未识别章节":
        _append_document_line(lines, section)
    for block in blocks:
        _append_document_line(lines, _block_document_text(block))
        if block.html:
            table_text = _html_to_text(block.html)
            _append_document_line(lines, _truncate_for_embedding(table_text))
    return "\n".join(lines).strip() or display_text


def _semantic_chunk_display_text(blocks: list[ParsedBlock], section: str) -> str:
    lines: list[str] = []
    if section != "未识别章节":
        _append_document_line(lines, section)
    for block in blocks:
        _append_document_line(lines, _block_document_text(block))
    return "\n".join(lines).strip()


def _block_document_text(block: ParsedBlock) -> str:
    text = block.text.strip()
    if block.image_path and text == block.image_path.strip():
        return ""
    if block.html and len(text) > 900:
        return "\n".join(_unique_strings([block.caption, _truncate_for_embedding(text, 600)]))
    return text


def _block_semantic_size(block: ParsedBlock) -> int:
    size = len(_block_document_text(block))
    if block.html:
        table_text = _html_to_text(block.html)
        if table_text and table_text not in block.text:
            size += min(len(table_text), 900)
    return size


def _block_section(block: ParsedBlock) -> str:
    return block.parent_section or block.section_hint or "未识别章节"


def _semantic_chunk_section(blocks: list[ParsedBlock]) -> str:
    for block in reversed(blocks):
        if block.type != "title" and _block_section(block) != "未识别章节":
            return _block_section(block)
    for block in reversed(blocks):
        if _block_section(block) != "未识别章节":
            return _block_section(block)
    return "未识别章节"


def _overlap_blocks(blocks: list[ParsedBlock]) -> list[ParsedBlock]:
    candidates = [block for block in blocks if block.type != "title" and _block_document_text(block)]
    return candidates[-SEMANTIC_CHUNK_OVERLAP_BLOCKS:]


def _evidence_window_blocks(
    core_blocks: list[ParsedBlock],
    all_blocks: list[ParsedBlock],
    block_positions: dict[str, int],
) -> list[ParsedBlock]:
    positions = [block_positions[block.block_id] for block in core_blocks if block.block_id in block_positions]
    if not positions:
        return core_blocks
    start = max(0, min(positions) - EVIDENCE_WINDOW_BLOCK_RADIUS)
    end = min(len(all_blocks) - 1, max(positions) + EVIDENCE_WINDOW_BLOCK_RADIUS)
    window = all_blocks[start : end + 1]
    section = _block_section(core_blocks[0])
    section_titles = [
        block
        for block in all_blocks[: min(positions)]
        if block.type == "title" and block.text and block.text in section
    ]
    return _unique_blocks([*section_titles[-3:], *window])


def _unique_blocks(blocks: list[ParsedBlock]) -> list[ParsedBlock]:
    seen: set[str] = set()
    result: list[ParsedBlock] = []
    for block in blocks:
        if block.block_id in seen:
            continue
        seen.add(block.block_id)
        result.append(block)
    return result


def _table_row_texts(block: ParsedBlock) -> list[str]:
    rows = _html_table_rows(block.html)
    if not rows:
        return []
    headers = rows[0] if _looks_like_header_row(rows[0]) else []
    data_rows = rows[1:] if headers else rows
    result: list[str] = []
    for row in data_rows:
        cells = [cell for cell in row if cell]
        if not cells:
            continue
        if headers and len(headers) == len(cells):
            result.append("；".join(f"{headers[index]}：{cell}" for index, cell in enumerate(cells)))
        else:
            result.append("；".join(cells))
    return _unique_strings(result)[:TABLE_ROW_MAX_ROWS_PER_TABLE]


def _html_table_rows(html: str) -> list[list[str]]:
    parser = _TableHTMLParser()
    parser.feed(html)
    return parser.rows


class _TableHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.rows: list[list[str]] = []
        self._current_row: list[str] | None = None
        self._current_cell: list[str] | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        normalized = tag.lower()
        if normalized == "tr":
            self._current_row = []
        elif normalized in {"td", "th"} and self._current_row is not None:
            self._current_cell = []

    def handle_data(self, data: str) -> None:
        if self._current_cell is not None:
            self._current_cell.append(data)

    def handle_endtag(self, tag: str) -> None:
        normalized = tag.lower()
        if normalized in {"td", "th"} and self._current_row is not None and self._current_cell is not None:
            cell = re.sub(r"\s+", " ", "".join(self._current_cell)).strip()
            self._current_row.append(cell)
            self._current_cell = None
        elif normalized == "tr" and self._current_row is not None:
            if any(cell.strip() for cell in self._current_row):
                self.rows.append(self._current_row)
            self._current_row = None


def _looks_like_header_row(row: list[str]) -> bool:
    joined = "".join(row)
    return bool(row) and any(keyword in joined for keyword in ["项目", "名称", "指标", "单位", "数量", "面积", "挖方", "填方", "位置", "结论"])


def _should_emit_table_row_chunks(block: ParsedBlock) -> bool:
    haystack = "\n".join([_block_section(block), block.caption, block.text, _html_to_text(block.html)])
    if any(keyword in haystack for keyword in TABLE_ROW_INCLUDE_KEYWORDS):
        return True
    if any(hint in haystack for hint in TABLE_ROW_EXCLUDE_HINTS):
        return False
    return False


def _caption_from_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines[:3]:
        if re.match(r"^(表|图)\s*[\d一二三四五六七八九十.-]+", line) or "表" in line[:8] or "图" in line[:8]:
            return line[:120]
    return ""


def _update_section_stack(section_stack: list[tuple[int, str]], title: str) -> list[tuple[int, str]]:
    return _replace_section_level(section_stack, _heading_level(title), title.strip())


def _is_section_title(title: str) -> bool:
    stripped = title.strip()
    if not stripped or len(stripped) > 80:
        return False
    if _section_hint(stripped):
        return True
    if re.match(r"^第[一二三四五六七八九十\d]+章", stripped):
        return True
    if re.match(r"^\d+(?:\.\d+)*\s+\S", stripped):
        return True
    if re.match(r"^[一二三四五六七八九十]+[、.．]", stripped):
        return True
    if re.match(r"^[（(][一二三四五六七八九十\d]+[）)]", stripped):
        return True
    return False


def _replace_section_level(section_stack: list[tuple[int, str]], level: int, title: str) -> list[tuple[int, str]]:
    if not title:
        return section_stack
    kept = [(item_level, item_title) for item_level, item_title in section_stack if item_level < level]
    return [*kept, (level, title)]


def _heading_level(title: str) -> int:
    stripped = title.strip()
    if re.match(r"^第[一二三四五六七八九十\d]+章", stripped):
        return 1
    match = re.match(r"^(\d+(?:\.\d+)*)", stripped)
    if match:
        return min(match.group(1).count(".") + 1, 5)
    if re.match(r"^[一二三四五六七八九十]+[、.．]", stripped):
        return 2
    if re.match(r"^[（(][一二三四五六七八九十\d]+[）)]", stripped):
        return 3
    return 2


def _section_path(section_stack: list[tuple[int, str]]) -> str:
    return " / ".join(title for _, title in section_stack)


def _append_document_line(lines: list[str], text: str) -> None:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    if not normalized:
        return
    if any(normalized == existing or normalized in existing for existing in lines):
        return
    lines.append(normalized)


def _truncate_for_embedding(text: str, limit: int = 900) -> str:
    normalized = re.sub(r"\s+", " ", text).strip()
    if len(normalized) <= limit:
        return normalized
    return normalized[:limit].rstrip() + "..."


def _number_list(values: Any) -> list[float]:
    if not isinstance(values, list):
        return []
    result: list[float] = []
    for value in values:
        try:
            result.append(round(float(value), 2))
        except (TypeError, ValueError):
            continue
    return result


def _optional_int(value: Any) -> int | None:
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _unique_strings(values: Any) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        text = str(value or "").strip()
        if text and text not in seen:
            seen.add(text)
            result.append(text)
    return result


def _html_to_text(html: str) -> str:
    text = re.sub(r"<[^>]+>", " ", html)
    return re.sub(r"\s+", " ", text).strip()


def _issues_from_configured_rules(
    session_id: str,
    chunks: list[ReviewChunk],
    full_text: str,
    fields: list[dict[str, Any]],
    rules: list[dict[str, Any]],
    limit: int = 12,
    evidence_store: Any | None = None,
) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    if limit <= 0:
        return issues
    for rule in rules:
        if not isinstance(rule, dict):
            continue
        structured_issue = _issue_from_structured_check_item(session_id, chunks, fields, rule, evidence_store)
        if structured_issue:
            issues.append(structured_issue)
            if len(issues) >= limit:
                break
            continue
        target_fields = [str(item) for item in rule.get("target_fields", []) if str(item).strip()]
        if not target_fields:
            continue
        matched = [field for field in target_fields if field in full_text]
        missing = [field for field in target_fields if field not in full_text]
        if not missing:
            continue
        # Keep the first version conservative: only emit rules with at least
        # one adjacent evidence hit, or rules whose target is a known high-frequency item.
        high_frequency = any(key in "、".join(target_fields) for key in ["土石方", "弃方", "表土", "占地", "防治责任范围", "监测", "投资"])
        if not matched and not high_frequency:
            continue
        evidence_matches = _evidence_matches_from_chunks(chunks, matched or target_fields)
        chunk = _chunk_for_evidence_match(chunks, evidence_matches[0]) if evidence_matches else _best_chunk(chunks, matched or target_fields)
        severity = _severity_from_policy(rule.get("severity_policy", ""))
        actual = "已命中：" + "、".join(matched) if matched else "材料中未见明确表述"
        if missing:
            actual += "；待核验：" + "、".join(missing[:6])
        issues.append(
            _issue(
                session_id=session_id,
                rule={
                    "rule_id": rule.get("rule_id", "WSB-CONFIG"),
                    "rule_name": rule.get("rule_name", "规则库审查"),
                    "rule_category": rule.get("category", "规则库审查"),
                    "severity": severity,
                    "expected": rule.get("evidence_requirement", "应满足规则库证据要求"),
                },
                issue_desc=_generic_rule_issue_desc(rule, matched, missing),
                evidence_text="",
                actual=actual,
                expected=rule.get("evidence_requirement", "应满足规则库证据要求"),
                chunk=chunk,
                risk_override=severity,
                confidence=68 if matched else 58,
                extra_reasoning=_generic_rule_reasoning(rule, matched, missing),
                evidence_matches=evidence_matches,
            )
        )
        if len(issues) >= limit:
            break
    return issues


def _issue_from_structured_check_item(
    session_id: str,
    chunks: list[ReviewChunk],
    fields: list[dict[str, Any]],
    rule: dict[str, Any],
    evidence_store: Any | None = None,
) -> dict[str, Any] | None:
    if not isinstance(rule.get("evidence_slots"), list) and not isinstance(rule.get("formula_checks"), list):
        return None

    from app.services.review_agent_service import build_evidence_slot_package
    from app.services.review_formula_service import execute_formula_checks
    from app.services.rag_service import RAGReviewError

    retrieval_trace = {
        "requested_store": evidence_store is not None,
        "used_store": evidence_store is not None,
        "degraded": False,
        "reason_type": "",
    }
    try:
        evidence_slot_package = build_evidence_slot_package(
            rule,
            chunks,
            evidence_store,
            use_bm25=True,
            use_neighbors=True,
            use_rerank=evidence_store is not None,
        )
    except RAGReviewError as exc:
        if evidence_store is None or not _is_vector_retrieval_failure(exc):
            raise
        retrieval_trace = {
            "requested_store": True,
            "used_store": False,
            "degraded": True,
            "reason_type": "vector_retrieval_failed",
        }
        evidence_slot_package = build_evidence_slot_package(
            rule,
            chunks,
            None,
            use_bm25=True,
            use_neighbors=True,
            use_rerank=False,
        )
    formula_check_results = execute_formula_checks(
        [item for item in rule.get("formula_checks", []) if isinstance(item, dict)],
        fields,
    )
    earthwork_audit_results = _earthwork_audit_results(rule, fields)
    from app.services.project_composition_service import analyze_project_composition_consistency

    project_comparison = analyze_project_composition_consistency(chunks, rule)
    project_comparison_status = (
        str(project_comparison.get("status") or "")
        if isinstance(project_comparison, dict)
        else ""
    )
    missing_slot_ids = [
        str(slot_id)
        for slot_id in evidence_slot_package.get("missing_required_slot_ids", [])
        if str(slot_id).strip()
    ]
    formula_checks = [item for item in formula_check_results.get("checks", []) if isinstance(item, dict)]
    failed_formula_ids = [
        str(item.get("formula_check_id") or "")
        for item in formula_checks
        if item.get("status") == "fail"
    ]
    missing_formula_ids = [
        str(item.get("formula_check_id") or "")
        for item in formula_checks
        if item.get("status") in {"missing", "unsupported"}
    ]
    missing_audit_ids = [
        str(item.get("audit_check_id") or "")
        for item in earthwork_audit_results.get("checks", [])
        if isinstance(item, dict) and item.get("status") == "missing"
    ]
    project_comparison_blocks = project_comparison_status in {"mismatch", "needs_review"}
    if (
        not missing_slot_ids
        and not failed_formula_ids
        and not missing_formula_ids
        and not missing_audit_ids
        and not project_comparison_blocks
    ):
        return None

    if missing_slot_ids or missing_formula_ids:
        status = "needs_evidence"
    elif failed_formula_ids or project_comparison_status == "mismatch":
        status = "issue"
    else:
        status = "needs_evidence" if missing_audit_ids or project_comparison_status == "needs_review" else "issue"
    evidence_matches = _evidence_matches_from_slot_package(evidence_slot_package, chunks)
    chunk = _chunk_for_evidence_match(chunks, evidence_matches[0]) if evidence_matches else _best_chunk(chunks, _structured_issue_keywords(rule, evidence_slot_package))
    page = chunk.page_range[0] if chunk else 1
    evidence = chunk.text[:500] if chunk else ""
    rule_name = str(rule.get("rule_name") or rule.get("review_sub_type") or "配置化审查项")
    category = str(rule.get("category") or rule.get("review_type") or "配置化审查")
    expected = str(rule.get("evidence_requirement") or rule.get("expected_result") or "应满足配置化审查要求")
    actual_value, finding_reason = _structured_issue_reason(
        missing_slot_ids,
        missing_formula_ids,
        failed_formula_ids,
        missing_audit_ids,
        project_comparison_status,
        project_comparison,
    )
    issue_id = str(uuid.uuid4())
    source_bbox_list = _source_bbox_list_from_matches(evidence_matches) or (chunk.bbox_list if chunk else [])
    evidence_nodes = _block_ids_from_matches(evidence_matches) or [item.get("block_id") for item in (chunk.bbox_list if chunk else [])]
    source_pages = _source_pages_from_matches(evidence_matches) or ([page] if page else [])
    reasoning_summary = _review_reasoning_summary(actual_value, expected, None, f"{rule_name}：{finding_reason}，需补齐后复核。")
    suggestion = "补齐配置项所需证据或修正文档数据后重新审查。"
    risk_level = _severity_from_policy(str(rule.get("severity_policy") or "")) if status == "issue" else "HIGH"
    reasoning = {
        "issue_type": category,
        "rule_id": str(rule.get("rule_id") or rule.get("id") or "WSB-CONFIG"),
        "rule_name": rule_name,
        "actual_value": actual_value,
        "expected_value": expected,
        "evidence_nodes": evidence_nodes,
        "source_pages": source_pages,
        "source_bbox_list": source_bbox_list,
        "evidence_slot_package": evidence_slot_package,
        "retrieval_trace": retrieval_trace,
        "formula_check_results": formula_check_results,
        "earthwork_audit_results": earthwork_audit_results,
        "project_composition_consistency": project_comparison,
        "review_status": status,
        "conclusion_type": status,
        "reasoning_summary": reasoning_summary,
        "review_result": _review_result(
            issue_id=issue_id,
            session_status="待审核",
            review_topic=category,
            review_item=str(rule.get("review_sub_type") or rule_name),
            rule_id=str(rule.get("rule_id") or rule.get("id") or "WSB-CONFIG"),
            rule_name=rule_name,
            risk_level=risk_level,
            issue_desc=f"{rule_name}：{finding_reason}，需补齐后复核。",
            evidence_matches=evidence_matches or _evidence_matches_from_chunk(chunk),
            source_bbox_list=source_bbox_list,
            source_pages=source_pages,
            reasoning_summary=reasoning_summary,
            fix_suggestion=suggestion,
            confidence=76,
        ),
    }
    return {
        "id": issue_id,
        "session_id": session_id,
        "clause_text": evidence,
        "page_number": page,
        "paragraph_index": 0,
        "highlight_anchor": chunk.chunk_id if chunk else f"page{page}",
        "char_offset_start": chunk.char_start if chunk else 0,
        "char_offset_end": chunk.char_end if chunk else len(evidence),
        "risk_level": risk_level,
        "confidence_score": 76,
        "source_type": "rule_engine",
        "risk_category": category,
        "ai_finding": f"{rule_name}：{finding_reason}，需补齐后复核。",
        "ai_reasoning": json.dumps(reasoning, ensure_ascii=False),
        "suggested_revision": suggestion,
        "human_decision": "pending",
    }


def _is_vector_retrieval_failure(exc: Exception) -> bool:
    return str(exc).startswith("vector retrieval failed:")


def _earthwork_audit_results(rule: dict[str, Any], fields: list[dict[str, Any]]) -> dict[str, Any]:
    if not _should_run_earthwork_audit(rule):
        return {"source": "earthwork_audit", "status": "not_applicable", "check_count": 0, "missing_count": 0, "checks": []}
    from app.services.earthwork_audit_service import execute_earthwork_audit

    return execute_earthwork_audit(fields)


def _should_run_earthwork_audit(rule: dict[str, Any]) -> bool:
    text = " ".join(
        str(value or "")
        for value in [
            rule.get("rule_id"),
            rule.get("rule_name"),
            rule.get("review_type"),
            rule.get("review_sub_type"),
            rule.get("category"),
            rule.get("review_criteria"),
            rule.get("expected_result"),
        ]
    )
    if "土石方" in text or "表土" in text:
        return True
    earthwork_fields = {"excavation_volume", "fill_volume", "borrow_volume", "spoil_volume"}
    for check in rule.get("formula_checks", []):
        if not isinstance(check, dict):
            continue
        configured = set(_string_list_for_structured_rule(check.get("left_fields"))) | set(_string_list_for_structured_rule(check.get("right_fields")))
        if configured.intersection(earthwork_fields):
            return True
    return False


def _string_list_for_structured_rule(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []


def _structured_issue_reason(
    missing_slot_ids: list[str],
    missing_formula_ids: list[str],
    failed_formula_ids: list[str],
    missing_audit_ids: list[str],
    project_comparison_status: str = "",
    project_comparison: dict[str, Any] | None = None,
) -> tuple[str, str]:
    if missing_slot_ids:
        return "缺失必填 evidence_slots：" + "、".join(missing_slot_ids), "必填证据槽位缺失"
    if missing_formula_ids:
        return "公式校验缺少必要字段或单位不支持：" + "、".join(item for item in missing_formula_ids if item), "公式校验证据缺失"
    if failed_formula_ids:
        return "公式校验未通过：" + "、".join(item for item in failed_formula_ids if item), "公式校验未通过"
    if missing_audit_ids:
        return "土石方结构化审计缺项：" + "、".join(item for item in missing_audit_ids if item), "土石方结构化审计缺项"
    if project_comparison_status == "mismatch":
        details = _project_comparison_detail_text(project_comparison)
        actual = "项目概要与立项或主体设计文件存在关键建设规模差异"
        return f"{actual}：{details}" if details else actual, "项目组成一致性不通过" + (f"：{details}" if details else "")
    if project_comparison_status == "needs_review":
        details = _project_comparison_detail_text(project_comparison)
        actual = "项目概要或立项/主体设计文件缺少可结构化比较字段"
        return f"{actual}：{details}" if details else actual, "项目组成一致性证据不足" + (f"：{details}" if details else "")
    return "配置化审查命中待复核条件", "配置化审查需复核"


def _project_comparison_detail_text(project_comparison: dict[str, Any] | None) -> str:
    if not isinstance(project_comparison, dict):
        return ""
    key_findings = project_comparison.get("key_findings")
    if isinstance(key_findings, list):
        details = [str(item).strip() for item in key_findings if str(item).strip()]
        if details:
            return "；".join(details[:3])
    reason = str(project_comparison.get("reason") or "").strip()
    return reason


def _generic_rule_issue_desc(rule: dict[str, Any], matched: list[str], missing: list[str]) -> str:
    rule_name = str(rule.get("rule_name") or "规则库审查")
    matched_text = "、".join(matched) if matched else "无"
    missing_text = "、".join(missing[:6]) if missing else "无"
    requirement = str(rule.get("evidence_requirement") or "应满足规则库证据要求")
    return (
        f"{rule_name}：按目标字段逐项核验，材料中已定位「{matched_text}」，"
        f"但未定位到「{missing_text}」。依据规则要求“{requirement}”，判定为证据不足，需要复核。"
    )


def _generic_rule_reasoning(rule: dict[str, Any], matched: list[str], missing: list[str]) -> dict[str, Any]:
    target_fields = [str(item) for item in rule.get("target_fields", []) if str(item).strip()]
    requirement = str(rule.get("evidence_requirement") or "应满足规则库证据要求")
    rule_source = str(rule.get("rule_source") or "")
    severity_policy = str(rule.get("severity_policy") or "")
    steps = [
        f"1. 读取规则目标字段：{'、'.join(target_fields) or '-'}。",
        f"2. 在当前审查对象全文中检索目标字段，已命中：{'、'.join(matched) if matched else '无'}。",
        f"3. 未命中字段按“材料中未见明确表述”处理，待核验：{'、'.join(missing) if missing else '无'}。",
        f"4. 规则要求：{requirement}",
        "5. 只要存在必需目标字段未形成可核验证据，就不判定通过，输出为证据不足/需复核。",
    ]
    if severity_policy:
        steps.append(f"6. 风险等级口径：{severity_policy}")
    return {
        "matched_target_fields": matched,
        "missing_target_fields": missing,
        "target_fields": target_fields,
        "rule_source": rule_source,
        "severity_policy": severity_policy,
        "evidence_requirement": requirement,
        "judgement_basis": (
            f"{rule_source + '；' if rule_source else ''}"
            f"规则要求：{requirement}；"
            "判定逻辑：target_fields 中的字段必须在方案正文、附表、附图或支撑材料中形成可核验证据；"
            "未命中的目标字段按材料中未见明确表述处理。"
        ),
        "judgement_steps": steps,
    }


def _evidence_matches_from_slot_package(
    evidence_slot_package: dict[str, Any],
    chunks: list[ReviewChunk],
    limit: int = 8,
) -> list[dict[str, Any]]:
    matches: list[dict[str, Any]] = []
    seen: set[str] = set()
    for slot in evidence_slot_package.get("slots", []):
        if not isinstance(slot, dict):
            continue
        for key in ["prompt_matches", "trace_matches", "matches"]:
            for match in slot.get(key, []):
                if not isinstance(match, dict):
                    continue
                chunk_id = str(match.get("chunk_id") or "")
                if chunk_id and chunk_id in seen:
                    continue
                seen.add(chunk_id)
                enriched = dict(match)
                enriched["slot_id"] = slot.get("slot_id")
                enriched["slot_label"] = slot.get("label")
                if not enriched.get("anchors"):
                    chunk = _chunk_for_evidence_match(chunks, enriched)
                    if chunk:
                        enriched.update(_evidence_match_from_chunk(chunk, []))
                matches.append(enriched)
                if len(matches) >= limit:
                    return matches
    return matches


def _evidence_matches_from_chunks(
    chunks: list[ReviewChunk],
    keywords: list[str],
    limit: int = 5,
) -> list[dict[str, Any]]:
    scored: list[tuple[int, int, ReviewChunk, list[str]]] = []
    normalized_keywords = [keyword for keyword in keywords if keyword]
    for index, chunk in enumerate(chunks):
        text = chunk.text or ""
        matched_terms = [keyword for keyword in normalized_keywords if keyword in text]
        if not matched_terms:
            continue
        scored.append((len(matched_terms), -index, chunk, matched_terms))
    scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return [_evidence_match_from_chunk(chunk, matched_terms) for _, _, chunk, matched_terms in scored[:limit]]


def _evidence_matches_from_chunk(chunk: ReviewChunk | None) -> list[dict[str, Any]]:
    return [_evidence_match_from_chunk(chunk, [])] if chunk else []


def _evidence_match_from_chunk(chunk: ReviewChunk, matched_terms: list[str]) -> dict[str, Any]:
    page_start = chunk.page_range[0] if chunk.page_range else None
    page_end = chunk.page_range[-1] if chunk.page_range else page_start
    anchors = chunk.bbox_list or []
    block_ids = [str(anchor.get("block_id")) for anchor in anchors if anchor.get("block_id")]
    return {
        "chunk_id": chunk.chunk_id,
        "page": page_start,
        "page_end": page_end,
        "primary_page": page_start,
        "page_range": [page_start, page_end] if page_start is not None and page_end is not None else [],
        "section": chunk.section,
        "anchors": anchors,
        "block_ids": block_ids,
        "bbox_count": len(anchors),
        "matched_terms": matched_terms,
        "retrieval_sources": ["keyword"],
        "text": chunk.text[:1600],
    }


def _chunk_for_evidence_match(chunks: list[ReviewChunk], match: dict[str, Any]) -> ReviewChunk | None:
    chunk_id = str(match.get("chunk_id") or "")
    if not chunk_id:
        return None
    return next((chunk for chunk in chunks if chunk.chunk_id == chunk_id), None)


def _source_bbox_list_from_matches(matches: list[dict[str, Any]]) -> list[dict[str, Any]]:
    anchors: list[dict[str, Any]] = []
    seen: set[str] = set()
    for match in matches:
        for anchor in match.get("anchors", []):
            if not isinstance(anchor, dict):
                continue
            key = f"{anchor.get('page')}:{anchor.get('block_id')}:{anchor.get('bbox')}"
            if key in seen:
                continue
            seen.add(key)
            anchors.append(anchor)
    return anchors


def _block_ids_from_matches(matches: list[dict[str, Any]]) -> list[str]:
    block_ids: list[str] = []
    seen: set[str] = set()
    for match in matches:
        candidates = match.get("block_ids")
        if not isinstance(candidates, list):
            candidates = [anchor.get("block_id") for anchor in match.get("anchors", []) if isinstance(anchor, dict)]
        for block_id in candidates:
            text = str(block_id or "").strip()
            if not text or text in seen:
                continue
            seen.add(text)
            block_ids.append(text)
    return block_ids


def _source_pages_from_matches(matches: list[dict[str, Any]]) -> list[int]:
    pages: set[int] = set()
    for match in matches:
        for key in ["page", "primary_page"]:
            value = match.get(key)
            if isinstance(value, int):
                pages.add(value)
        page_range = match.get("page_range")
        if isinstance(page_range, list):
            pages.update(int(page) for page in page_range if isinstance(page, int))
    return sorted(page for page in pages if page > 0)


def _evidence_text_from_matches(matches: list[dict[str, Any]], limit: int = 5) -> str:
    lines: list[str] = []
    for index, match in enumerate(matches[:limit], start=1):
        page = match.get("page") or match.get("primary_page") or "-"
        chunk_id = str(match.get("chunk_id") or "-")
        section = str(match.get("section") or "").strip()
        text = re.sub(r"\s+", " ", str(match.get("text") or "")).strip()
        if not text:
            continue
        prefix = f"{index}. {chunk_id} p.{page}"
        if section:
            prefix += f" {section}"
        lines.append(f"{prefix}：{text[:320]}")
    return "\n".join(lines)


def _review_reasoning_summary(
    actual: str,
    expected: str,
    extra_reasoning: dict[str, Any] | None,
    issue_desc: str,
) -> str:
    steps = []
    if isinstance(extra_reasoning, dict):
        steps = [str(item).strip() for item in extra_reasoning.get("judgement_steps", []) if str(item).strip()]
    if steps:
        return "；".join(steps[:6])
    return "；".join(
        item
        for item in [
            f"判定对象：{issue_desc}",
            f"实际命中：{actual}",
            f"规则要求：{expected}",
            "判定结论：存在未命中或待核验证据时，不判定为通过。",
        ]
        if item
    )


def _review_result(
    *,
    issue_id: str,
    session_status: str,
    review_topic: str,
    review_item: str,
    rule_id: str,
    rule_name: str,
    risk_level: str,
    issue_desc: str,
    evidence_matches: list[dict[str, Any]],
    source_bbox_list: list[dict[str, Any]],
    source_pages: list[int],
    reasoning_summary: str,
    fix_suggestion: str,
    confidence: int,
) -> dict[str, Any]:
    return {
        "issue_id": issue_id,
        "review_topic": review_topic,
        "review_item": review_item,
        "rule_id": rule_id,
        "rule_name": rule_name,
        "risk_level": risk_level,
        "issue_desc": issue_desc,
        "evidence_text": _evidence_text_from_matches(evidence_matches),
        "evidence_nodes": [
            {
                "chunk_id": match.get("chunk_id"),
                "page": match.get("page") or match.get("primary_page"),
                "page_end": match.get("page_end"),
                "section": match.get("section"),
                "block_ids": match.get("block_ids") or [],
                "bbox_count": match.get("bbox_count") or len(match.get("anchors", [])),
                "matched_terms": match.get("matched_terms") or [],
                "retrieval_sources": match.get("retrieval_sources") or [],
                "text": str(match.get("text") or "")[:800],
            }
            for match in evidence_matches
        ],
        "source_pages": source_pages,
        "source_bbox_list": source_bbox_list,
        "reasoning_summary": reasoning_summary,
        "fix_suggestion": fix_suggestion,
        "confidence": confidence,
        "review_status": session_status,
    }


def _structured_issue_keywords(rule: dict[str, Any], evidence_slot_package: dict[str, Any]) -> list[str]:
    keywords: list[str] = []
    for slot in evidence_slot_package.get("slots", []):
        if not isinstance(slot, dict):
            continue
        keywords.extend(str(term) for term in slot.get("matched_expected_terms", []) if str(term).strip())
        keywords.extend(str(term) for term in slot.get("expected_terms", []) if str(term).strip())
    keywords.extend(str(item) for item in rule.get("target_fields", []) if str(item).strip())
    return keywords or [str(rule.get("rule_name") or rule.get("review_sub_type") or "")]


def _severity_from_policy(policy: str) -> str:
    if "重大" in policy or "严重" in policy:
        return "HIGH"
    if "一般" in policy:
        return "MEDIUM"
    return "LOW"


def _parse_pdf(file_path: str) -> list[ParsedBlock]:
    import fitz

    blocks: list[ParsedBlock] = []
    cursor = 0
    with fitz.open(file_path) as doc:
        for page_index, page in enumerate(doc, start=1):
            data = page.get_text("dict")
            for raw in data.get("blocks", []):
                lines = []
                for line in raw.get("lines", []):
                    line_text = "".join(
                        span.get("text", "")
                        for span in line.get("spans", [])
                    ).strip()
                    if line_text:
                        lines.append(line_text)
                text = "\n".join(lines).strip()
                if not text:
                    continue
                start = cursor
                cursor += len(text) + 1
                blocks.append(
                    ParsedBlock(
                        block_id=f"p{page_index}-b{len(blocks) + 1:05d}",
                        page=page_index,
                        bbox=[round(float(v), 2) for v in raw.get("bbox", [])],
                        text=text,
                        type=_classify_block(text),
                        section_hint=_section_hint(text),
                        char_start=start,
                        char_end=start + len(text),
                    )
                )
    return blocks


def _parse_docx(file_path: str) -> list[ParsedBlock]:
    from docx import Document

    blocks: list[ParsedBlock] = []
    cursor = 0
    doc = Document(file_path)
    for index, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()
        if not text:
            continue
        start = cursor
        cursor += len(text) + 1
        blocks.append(
            ParsedBlock(
                block_id=f"docx-b{index:05d}",
                page=1,
                bbox=[],
                text=text,
                type=_classify_block(text),
                section_hint=_section_hint(text),
                char_start=start,
                char_end=start + len(text),
            )
        )
    for table_index, table in enumerate(doc.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        text = "\n".join(row for row in rows if row.strip())
        if not text:
            continue
        start = cursor
        cursor += len(text) + 1
        blocks.append(
            ParsedBlock(
                block_id=f"docx-t{table_index:05d}",
                page=1,
                bbox=[],
                text=text,
                type="table",
                section_hint="表格",
                char_start=start,
                char_end=start + len(text),
            )
        )
    return blocks


def _classify_block(text: str) -> str:
    stripped = text.strip()
    if len(stripped) <= 40 and re.match(r"^([一二三四五六七八九十]+[、.．]|第[一二三四五六七八九十\d]+章|\d+(\.\d+)*\s*)", stripped):
        return "title"
    if "|" in stripped or re.search(r"\b(合计|单位|面积|数量)\b", stripped):
        return "table" if len(stripped) > 60 else "paragraph"
    return "paragraph"


def _section_hint(text: str) -> str:
    for section, keywords in SECTION_KEYWORDS.items():
        if any(keyword in text for keyword in keywords):
            return section
    return ""


def _detect_sections(text: str) -> list[str]:
    return [section for section, keywords in SECTION_KEYWORDS.items() if any(keyword in text for keyword in keywords)]


def _field(name: str, match: dict[str, Any] | None, chunks: list[ReviewChunk]) -> dict[str, Any]:
    if not match:
        return {
            "field_name": name,
            "value": "",
            "normalized_value": "",
            "source_span": None,
            "section": "",
            "confidence": 35,
        }
    chunk = _chunk_for_span(chunks, match["start"])
    return {
        "field_name": name,
        "value": match["value"],
        "normalized_value": match.get("normalized_value", match["value"]),
        "source_span": {"char_start": match["start"], "char_end": match["end"]},
        "section": chunk.section if chunk else "",
        "confidence": match.get("confidence", 78),
    }


def _match_after(text: str, patterns: list[str]) -> dict[str, Any] | None:
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            value = match.group(1).strip()
            return {"value": value, "start": match.start(1), "end": match.end(1), "confidence": 82}
    return None


def _match_metric(text: str, labels: list[str]) -> dict[str, Any] | None:
    label_pattern = "|".join(re.escape(label) for label in labels)
    pattern = rf"({label_pattern})[^\d一二三四五六七八九十零〇百千万亿]{{0,12}}(\d+(?:\.\d+)?)\s*([万亿]?(?:m3|m³|立方米|万m3|万m³|hm2|hm²|公顷|亩|万元|元)?)"
    match = re.search(pattern, text, re.IGNORECASE)
    if not match:
        return None
    value = (match.group(2) + (match.group(3) or "")).strip()
    return {
        "value": value,
        "normalized_value": match.group(2),
        "start": match.start(2),
        "end": match.end(3),
        "confidence": 76,
    }


def _match_keyword(text: str, keywords: list[str]) -> dict[str, Any] | None:
    for keyword in keywords:
        idx = text.find(keyword)
        if idx >= 0:
            return {"value": keyword, "start": idx, "end": idx + len(keyword), "confidence": 70}
    return None


def _chunk_for_span(chunks: list[ReviewChunk], offset: int) -> ReviewChunk | None:
    for chunk in chunks:
        if chunk.char_start <= offset <= chunk.char_end:
            return chunk
    return chunks[0] if chunks else None


def _best_chunk(chunks: list[ReviewChunk], keywords: list[str]) -> ReviewChunk | None:
    for keyword in keywords:
        for chunk in chunks:
            if keyword in chunk.text or keyword in chunk.section:
                return chunk
    return chunks[0] if chunks else None


def _issue(
    session_id: str,
    rule: dict[str, Any],
    issue_desc: str,
    evidence_text: str,
    actual: str,
    expected: str,
    chunk: ReviewChunk | None,
    risk_override: str | None = None,
    confidence: int = 78,
    extra_reasoning: dict[str, Any] | None = None,
    evidence_matches: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    issue_id = str(uuid.uuid4())
    page = chunk.page_range[0] if chunk else 1
    evidence = evidence_text or (chunk.text[:500] if chunk else "")
    matches = evidence_matches or _evidence_matches_from_chunk(chunk)
    source_bbox_list = _source_bbox_list_from_matches(matches) or (chunk.bbox_list if chunk else [])
    evidence_nodes = _block_ids_from_matches(matches) or [item.get("block_id") for item in (chunk.bbox_list if chunk else []) if item.get("block_id")]
    source_pages = _source_pages_from_matches(matches) or ([page] if page else [])
    risk_level = risk_override or rule["severity"]
    review_status = "待审核"
    fix_suggestion = _suggestion_for(rule)
    reasoning_summary = _review_reasoning_summary(actual, expected, extra_reasoning, issue_desc)
    reasoning = {
        "issue_type": rule["rule_category"],
        "rule_id": rule["rule_id"],
        "rule_name": rule["rule_name"],
        "actual_value": actual,
        "expected_value": expected,
        "evidence_nodes": evidence_nodes,
        "source_pages": source_pages,
        "source_bbox_list": source_bbox_list,
        "review_status": "pending",
        "conclusion_type": "issue" if risk_level != "LOW" else "attention",
        "reasoning_summary": reasoning_summary,
    }
    if isinstance(extra_reasoning, dict):
        reasoning.update(extra_reasoning)
    reasoning["review_result"] = _review_result(
        issue_id=issue_id,
        session_status=review_status,
        review_topic=str(rule.get("rule_category") or ""),
        review_item=str(rule.get("rule_name") or ""),
        rule_id=str(rule.get("rule_id") or ""),
        rule_name=str(rule.get("rule_name") or ""),
        risk_level=str(risk_level),
        issue_desc=issue_desc,
        evidence_matches=matches,
        source_bbox_list=source_bbox_list,
        source_pages=source_pages,
        reasoning_summary=reasoning_summary,
        fix_suggestion=fix_suggestion,
        confidence=confidence,
    )
    return {
        "id": issue_id,
        "session_id": session_id,
        "clause_text": evidence,
        "page_number": page,
        "paragraph_index": 0,
        "highlight_anchor": chunk.chunk_id if chunk else f"page{page}",
        "char_offset_start": chunk.char_start if chunk else 0,
        "char_offset_end": chunk.char_end if chunk else len(evidence),
        "risk_level": risk_level,
        "confidence_score": confidence,
        "source_type": "rule_engine",
        "risk_category": rule["rule_category"],
        "ai_finding": issue_desc,
        "ai_reasoning": json.dumps(reasoning, ensure_ascii=False),
        "suggested_revision": fix_suggestion,
        "human_decision": "pending",
    }


def _suggestion_for(rule: dict[str, Any]) -> str:
    suggestions = {
        "SWC-FORM-001": "补充缺失章节，或在目录及正文中明确对应章节名称与内容。",
        "SWC-FIELD-001": "在项目概况中补充项目名称、建设单位、建设地点、建设性质等基础信息。",
        "SWC-TECH-001": "补充扰动地表面积、防治责任范围面积及对应计算依据。",
        "SWC-TECH-002": "复核土石方平衡表，统一挖方、填方、借方、弃方口径并说明去向。",
        "SWC-TECH-003": "补充表土剥离、临时保存、防护和回覆利用措施。",
        "SWC-TECH-004": "补充弃渣场位置、容量、拦挡排水措施，或说明弃方合法去向。",
        "SWC-ATTR-001": "补充项目区水土流失重点防治区属性及适用规范依据。",
    }
    return suggestions.get(rule["rule_id"], "请补充材料并由审查人员复核。")


def _to_number(value: str) -> float | None:
    if not value:
        return None
    try:
        return float(value)
    except ValueError:
        return None


def _write_json(path: Path, data: Any) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
